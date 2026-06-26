"""Поиск ФИО из файлов через Sauron.

Формат отчёта: два листа XLSX
  «Итог по людям»   — одна строка = один исходный человек из файла
  «Связанные лица»  — одна строка = одно связанное лицо (child record)

Для связанных лиц:
  • Только mobile-номера, максимум TOP_PHONES_PER_RELATED лучших
  • Приоритет: свежие источники (2024-2026), Maxim/taxsee, многократное
    подтверждение, мобильный префикс
  • Отдельные колонки: ВК, Одноклассники, другие соцсети
  • Парсинг даты рождения и типа связи из поля «Связь с лицом»
  • Детектор Макса/taxsee по источникам Sauron
  • Глобальный dedup телефонов

Форматы файлов: txt, csv, xlsx, xls, docx, pdf.

Env:
  SAURON_FILE_MAX_FIO         — макс. ФИО за файл      (умолч. 30)
  SAURON_FILE_DELAY_SEC       — пауза сек              (умолч. 2)
  SAURON_FILE_MAX_RELATED     — макс. связанных на ФИО (умолч. 3)
  SAURON_FILE_TOP_REL_PHONES  — топ тел. на связанного (умолч. 4)
"""
from __future__ import annotations

import os
import re
import io
import csv
import time
import logging
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)

# ── Лимиты ───────────────────────────────────────────────────────────────────
DEFAULT_MAX_FIO          = int(os.environ.get("SAURON_FILE_MAX_FIO",         "30"))
DEFAULT_DELAY_SEC        = float(os.environ.get("SAURON_FILE_DELAY_SEC",     "2"))
DEFAULT_MAX_RELATED      = int(os.environ.get("SAURON_FILE_MAX_RELATED",     "3"))
TOP_PHONES_PER_RELATED   = int(os.environ.get("SAURON_FILE_TOP_REL_PHONES",  "4"))

# Обратная совместимость
DEFAULT_MAX_RECORDS = DEFAULT_MAX_FIO

# ── Опциональные библиотеки ───────────────────────────────────────────────────
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    _HAS_OPENPYXL = True
except ImportError:
    _HAS_OPENPYXL = False

try:
    import xlrd
    _HAS_XLRD = True
except ImportError:
    _HAS_XLRD = False

try:
    from docx import Document as DocxDocument
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    import pdfplumber
    _HAS_PDF = True
except ImportError:
    _HAS_PDF = False


# ═════════════════════════════════════════════════════════════════════════════
# Структуры данных
# ═════════════════════════════════════════════════════════════════════════════

@dataclass
class RelatedPerson:
    """Одно связанное лицо."""
    fio:            str = ""
    dob:            str = ""          # дата рождения
    relation:       str = ""          # тип связи / причина
    address:        str = ""
    liquid_phones:  list[str] = field(default_factory=list)
    discarded_cnt:  int = 0
    phone_note:     str = ""
    in_maxim:       bool = False
    maxim_source:   str = ""
    vk:             str = ""
    ok:             str = ""
    other_social:   str = ""
    sources:        str = ""
    comment:        str = ""
    raw_count:      int = 0


@dataclass
class PersonResult:
    """Результат поиска одного ФИО из файла."""
    row_num:        int = 0
    source_fio:     str = ""
    found:          bool = False
    found_fio:      str = ""
    liquid_phones:  str = ""          # ликвидные тел. основного
    phones_raw:     str = ""
    discarded_phones: int = 0
    emails:         str = ""
    addresses:      str = ""
    connections:    str = ""
    in_maxim:       bool = False
    maxim_source:   str = ""
    vk:             str = ""
    ok:             str = ""
    other_social:   str = ""
    sources:        str = ""
    raw_count:      int = 0
    confidence:     str = ""
    phone_note:     str = ""
    error:          str = ""
    related: list[RelatedPerson] = field(default_factory=list)


# ═════════════════════════════════════════════════════════════════════════════
# Стоп-слова ФИО
# ═════════════════════════════════════════════════════════════════════════════

_NAME_STOPWORDS: frozenset[str] = frozenset({
    "Республика", "Область", "Край", "Округ", "Район", "Автономная",
    "Федерация", "Федеральный", "Федеральная", "Муниципальный", "Муниципальная",
    "Поселение", "Поселок", "Деревня", "Хутор", "Аул", "Станица",
    "Дагестан", "Башкортостан", "Татарстан", "Чечня", "Ингушетия",
    "Калмыкия", "Черкесия", "Адыгея", "Мордовия", "Удмуртия",
    "Чувашия", "Якутия", "Бурятия", "Хакасия", "Тыва", "Тува",
    "Карачаево", "Кабардино", "Балкария", "Коми", "Крым", "Алтай", "Марий",
    "Чувашская", "Чеченская", "Дагестанская", "Башкирская",
    "Татарская", "Ингушская", "Кабардинская", "Балкарская",
    "Калмыцкая", "Карачаевская", "Черкесская", "Адыгейская",
    "Мордовская", "Удмуртская", "Бурятская", "Тувинская",
    "Хакасская", "Якутская", "Алтайская",
    "Московская", "Ленинградская", "Нижегородская", "Самарская",
    "Саратовская", "Волгоградская", "Ростовская", "Краснодарская",
    "Ставропольская", "Свердловская", "Челябинская", "Тюменская",
    "Омская", "Новосибирская", "Иркутская", "Красноярская",
    "Хабаровская", "Воронежская", "Кемеровская", "Пермская",
    "Оренбургская", "Тульская", "Рязанская", "Ярославская",
    "Владимирская", "Ивановская", "Костромская", "Смоленская",
    "Тверская", "Брянская", "Орловская", "Курская", "Белгородская",
    "Липецкая", "Тамбовская", "Пензенская", "Ульяновская",
    "Кировская", "Вологодская", "Архангельская", "Мурманская",
    "Псковская", "Астраханская", "Калужская", "Калининградская",
    "Российская", "Советская", "Украинская", "Белорусская",
    "Улица", "Проспект", "Бульвар", "Площадь", "Переулок",
    "Набережная", "Шоссе", "Тракт", "Квартал", "Микрорайон",
    "Город", "Москва", "Санкт-Петербург",
    "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
    "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь",
    "Понедельник", "Вторник", "Среда", "Четверг", "Пятница",
    "Суббота", "Воскресенье",
    "Украина", "Харків", "Одеса", "Дніпро",
})

_GEO_ADJ_RE = re.compile(r'^[А-ЯЁ][а-яё]{3,}(ская|ский|ское|ской|ских|ском)$')


def _is_valid_fio_word(word: str) -> bool:
    if word in _NAME_STOPWORDS:
        return False
    if _GEO_ADJ_RE.match(word):
        return False
    return 2 <= len(word) <= 30


_FIO3_RE = re.compile(
    r'\b([А-ЯЁІЇЄ][а-яёіїє\'\-]{1,25})\s+'
    r'([А-ЯЁІЇЄ][а-яёіїє\'\-]{1,20})\s+'
    r'([А-ЯЁІЇЄ][а-яёіїє\'\-]{1,20})\b'
)
_FIO2_RE = re.compile(
    r'\b([А-ЯЁІЇЄ][а-яёіїє\'\-]{1,25})\s+'
    r'([А-ЯЁІЇЄ][а-яёіїє\'\-]{1,20})\b'
)


def extract_names(text: str) -> list[str]:
    found: list[str] = []
    seen: set[str] = set()
    for m in _FIO3_RE.finditer(text):
        parts = [m.group(1), m.group(2), m.group(3)]
        if not all(_is_valid_fio_word(p) for p in parts):
            continue
        name = ' '.join(parts)
        if name not in seen:
            seen.add(name)
            found.append(name)
    for m in _FIO2_RE.finditer(text):
        p1, p2 = m.group(1), m.group(2)
        if not _is_valid_fio_word(p1) or not _is_valid_fio_word(p2):
            continue
        name = f'{p1} {p2}'
        if name not in seen and not any(name in ex for ex in seen):
            seen.add(name)
            found.append(name)
    return found


# ═════════════════════════════════════════════════════════════════════════════
# Детектор колонок
# ═════════════════════════════════════════════════════════════════════════════

_FIO_COLS    = frozenset({'фио', 'fullname', 'full_name', 'ф.и.о', 'ф.и.о.', 'полное имя', 'наименование', 'person', 'человек'})
_LAST_COLS   = frozenset({'фамилия', 'lastname', 'last_name', 'surname', 'фам', 'last'})
_FIRST_COLS  = frozenset({'имя', 'firstname', 'first_name', 'name', 'first'})
_PATRON_COLS = frozenset({'отчество', 'patronymic', 'patronym', 'middlename', 'middle_name', 'отч', 'middle'})
_PERSON_COLS = frozenset({'клиент', 'client', 'абонент', 'subscriber', 'пассажир', 'passenger',
                           'сотрудник', 'employee', 'владелец', 'owner', 'заказчик'})


def _norm_header(s: str) -> str:
    return re.sub(r'[\s_\-\.]+', '', str(s).lower().strip())


def _detect_fio_cols(headers: list[str]) -> dict[str, int]:
    result = {'fio': -1, 'last': -1, 'first': -1, 'patron': -1}
    for i, h in enumerate(headers):
        n = _norm_header(h)
        if n in _FIO_COLS or n in _PERSON_COLS:
            result['fio'] = i
        elif n in _LAST_COLS and result['last'] == -1:
            result['last'] = i
        elif n in _FIRST_COLS and result['first'] == -1:
            result['first'] = i
        elif n in _PATRON_COLS and result['patron'] == -1:
            result['patron'] = i
    return result


def _build_fio_from_cols(row: list[str], col_map: dict[str, int]) -> str:
    def get(key: str) -> str:
        idx = col_map.get(key, -1)
        return str(row[idx]).strip() if 0 <= idx < len(row) else ''

    if col_map.get('fio', -1) >= 0 and get('fio'):
        return get('fio')
    return ' '.join(p for p in [get('last'), get('first'), get('patron')] if p)


# ═════════════════════════════════════════════════════════════════════════════
# Парсинг файлов
# ═════════════════════════════════════════════════════════════════════════════

def _decode(data: bytes) -> str:
    for enc in ('utf-8', 'cp1251', 'utf-16', 'latin-1'):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode('utf-8', errors='replace')


def _parse_csv_structured(data: bytes) -> tuple[list[str], list[list[str]]]:
    text = _decode(data)
    for sep in (',', ';', '\t', '|'):
        try:
            reader = csv.reader(io.StringIO(text), delimiter=sep)
            all_rows = [r for r in reader if any(c.strip() for c in r)]
            if len(all_rows) >= 2 and len(all_rows[0]) >= 2:
                return all_rows[0], all_rows[1:]
        except Exception:
            continue
    lines = [l for l in text.splitlines() if l.strip()]
    return [], [[l] for l in lines]


def _parse_xlsx_structured(data: bytes) -> tuple[list[str], list[list[str]]]:
    if not _HAS_OPENPYXL:
        return [], []
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        ws = wb.active
        rows = [[str(c.value or '').strip() for c in row] for row in ws.iter_rows()]
        wb.close()
        return (rows[0], rows[1:]) if len(rows) >= 2 else ([], rows)
    except Exception as e:
        logger.warning(f"xlsx parse error: {e}")
        return [], []


def _parse_xls_structured(data: bytes) -> tuple[list[str], list[list[str]]]:
    if not _HAS_XLRD:
        return [], []
    try:
        wb = xlrd.open_workbook(file_contents=data)
        ws = wb.sheet_by_index(0)
        rows = [[str(ws.cell_value(r, c)).strip() for c in range(ws.ncols)] for r in range(ws.nrows)]
        return (rows[0], rows[1:]) if len(rows) >= 2 else ([], rows)
    except Exception as e:
        logger.warning(f"xls parse error: {e}")
        return [], []


def _parse_docx_text(data: bytes) -> str:
    if not _HAS_DOCX:
        return ''
    lines = []
    try:
        doc = DocxDocument(io.BytesIO(data))
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if parts:
                    lines.append('\t'.join(parts))
    except Exception as e:
        logger.warning(f"docx parse error: {e}")
    return '\n'.join(lines)


def _parse_pdf_text(data: bytes) -> str:
    if not _HAS_PDF:
        return ''
    lines = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages[:20]:
                text = page.extract_text()
                if text:
                    lines.append(text)
    except Exception as e:
        logger.warning(f"pdf parse error: {e}")
    return '\n'.join(lines)


def parse_file(data: bytes, filename: str) -> tuple[list[dict], str]:
    fname = filename.lower()
    records: list[dict] = []
    seen_fio: set[str] = set()

    def _add(row_num: int, source_line: str, fio: str):
        key = fio.strip().lower()
        if not key or key in seen_fio:
            return
        seen_fio.add(key)
        records.append({'row_num': row_num, 'source_line': source_line, 'fio': fio.strip()})

    def _from_structured(headers: list[str], rows: list[list[str]]):
        col_map = _detect_fio_cols(headers)
        has_struct = any(v >= 0 for v in col_map.values())
        if has_struct:
            for i, row in enumerate(rows, 2):
                fio = _build_fio_from_cols(row, col_map)
                if fio:
                    _add(i, ' | '.join(str(c) for c in row if str(c).strip()), fio)
        else:
            for i, row in enumerate(rows, 2):
                line = ' '.join(str(c) for c in row if str(c).strip())
                for name in extract_names(line):
                    _add(i, line, name)

    if fname.endswith('.csv'):
        _from_structured(*_parse_csv_structured(data))
    elif fname.endswith('.xlsx'):
        if not _HAS_OPENPYXL:
            return [], "xlsx не поддерживается — openpyxl не установлен."
        _from_structured(*_parse_xlsx_structured(data))
    elif fname.endswith('.xls'):
        if not _HAS_XLRD:
            return [], "xls не поддерживается — xlrd не установлен."
        _from_structured(*_parse_xls_structured(data))
    elif fname.endswith('.txt'):
        text = _decode(data)
        for i, line in enumerate(text.splitlines(), 1):
            stripped = line.strip()
            if not stripped:
                continue
            words = stripped.split()
            if 2 <= len(words) <= 4:
                names = extract_names(stripped)
                if names:
                    for name in names:
                        _add(i, stripped, name)
                    continue
            for name in extract_names(stripped):
                _add(i, stripped, name)
    elif fname.endswith('.docx'):
        if not _HAS_DOCX:
            return [], "docx не поддерживается — python-docx не установлен."
        text = _parse_docx_text(data)
        for i, line in enumerate(text.splitlines(), 1):
            for name in extract_names(line.strip()):
                _add(i, line.strip(), name)
    elif fname.endswith('.pdf'):
        if not _HAS_PDF:
            return [], "pdf не поддерживается — pdfplumber не установлен."
        text = _parse_pdf_text(data)
        for i, line in enumerate(text.splitlines(), 1):
            for name in extract_names(line.strip()):
                _add(i, line.strip(), name)
    else:
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else '?'
        return [], f"Формат .{ext} не поддерживается.\nПоддерживаю: txt, csv, xlsx, xls, docx, pdf"

    if not records:
        return [], (
            "В файле не найдено ни одного ФИО.\n\n"
            "Убедись, что:\n"
            "• Для csv/xlsx — есть колонка «ФИО», «Фамилия» или «Имя»\n"
            "• Для txt — каждая строка содержит 2–3 слова с заглавной буквы (кириллица)"
        )
    return records, ''


# ═════════════════════════════════════════════════════════════════════════════
# Проверка ликвидности телефонов
# ═════════════════════════════════════════════════════════════════════════════

_ALL_SAME_RE = re.compile(r'^(\d)\1{9,}$')
_SEQUENTIAL  = {'1234567890', '0987654321', '9876543210', '0123456789'}
_TECH_PFXS   = ('70000', '79000', '700000', '710000', '71234', '70123')

# Мобильные RU-коды (после 79)
_RU_MOBILE_CODES = frozenset({
    '901','902','903','904','905','906','908','909',
    '910','911','912','913','914','915','916','917','918','919',
    '920','921','922','923','924','925','926','927','928','929',
    '930','931','932','933','934','936','937','938','939',
    '950','951','952','953','955','958','960','961','962','963',
    '964','965','966','967','968','969',
    '970','977','978','980','981','982','983','984','985','986','987','988','989',
    '990','991','992','993','994','995','996','997','998','999',
})


def normalize_phone(raw: str) -> Optional[str]:
    """Нормализует телефон. Возвращает None если невалидный/маска."""
    if not raw:
        return None
    digits = re.sub(r'\D', '', raw)
    if not digits:
        return None
    if not (7 <= len(digits) <= 15):
        return None

    # Нормализация
    if len(digits) == 11 and digits[0] in ('7', '8'):
        norm = '7' + digits[1:]
    elif len(digits) == 10 and digits[0] == '0':
        norm = '380' + digits[1:]
    elif len(digits) == 12 and digits.startswith('380'):
        norm = digits
    else:
        norm = digits

    # Маски
    if _ALL_SAME_RE.match(norm):
        return None
    if norm[-10:] in _SEQUENTIAL:
        return None
    if any(norm.startswith(p) for p in _TECH_PFXS):
        return None
    if norm.replace('0', '') == '':
        return None

    # RU: цифра после кода страны
    if len(norm) == 11 and norm.startswith('7'):
        second = norm[1]
        if second not in ('9', '4', '8', '3', '2'):
            return None
    return norm


def _is_mobile(norm: str) -> bool:
    """True если номер мобильный (RU 79xxx или UA 380xxx)."""
    if norm.startswith('79') and len(norm) == 11:
        code3 = norm[1:4]
        return code3 in _RU_MOBILE_CODES
    if norm.startswith('380') and len(norm) == 12:
        prefix = norm[3]
        return prefix in ('6', '7', '9', '5', '4', '3')
    return False


def _phone_score(raw: str, sources_ctx: str, freq: int = 1) -> float:
    """
    Скоринг телефона: выше → надёжнее.
    Факторы: мобильный, свежий источник, Maxim/taxsee, частота.
    """
    norm = normalize_phone(raw)
    if norm is None:
        return -1.0
    score = 0.0
    # Мобильный приоритет
    if _is_mobile(norm):
        score += 10.0
    else:
        score += 2.0   # городские тоже можно, но ниже приоритет
    # Свежесть источников (2024-2026)
    for yr in ('2026', '2025', '2024'):
        if yr in sources_ctx:
            score += 5.0
            break
    for yr in ('2023', '2022'):
        if yr in sources_ctx:
            score += 2.0
            break
    # Maxim/taxsee
    if re.search(r'taxsee|maxim|максим\s*такс', sources_ctx, re.I):
        score += 8.0
    # Частота встречи в записях
    score += min(freq, 5) * 1.5
    return score


# ═════════════════════════════════════════════════════════════════════════════
# Детектор Макса / taxsee
# ═════════════════════════════════════════════════════════════════════════════

_MAXIM_RE = re.compile(
    r'taxsee|'
    r'\bmaxim(um)?\b|'
    r'максим(а|у|ом|е|овск)?\b|'
    r'mac[s]?tax|мак[сш]такс|maxtax|'
    r'клиент[ыа]?\s*(maxim|макс|такси)|'
    r'база\s*(maxim|макс)|'
    r'driver\s*base',
    re.IGNORECASE,
)
_MAXIM_BROAD_RE = re.compile(
    r'\bмакс\b.{0,20}(такси|taxi|водит|шофер|driver)',
    re.IGNORECASE,
)


def check_maxim(sources: str, raw_records: Optional[list[dict]] = None) -> tuple[bool, str]:
    combined = sources or ''
    if raw_records:
        for rec in raw_records:
            combined += '\n' + str(rec.get('Источник', '') or rec.get('База', '') or '')

    m = _MAXIM_RE.search(combined)
    if m:
        start = max(0, m.start() - 15)
        end   = min(len(combined), m.end() + 15)
        return True, combined[start:end].replace('\n', ' ').strip()

    m2 = _MAXIM_BROAD_RE.search(combined)
    if m2:
        start = max(0, m2.start() - 10)
        end   = min(len(combined), m2.end() + 10)
        return True, combined[start:end].replace('\n', ' ').strip()

    return False, ''


# ═════════════════════════════════════════════════════════════════════════════
# Извлечение соцсетей
# ═════════════════════════════════════════════════════════════════════════════

_VK_RE  = re.compile(r'vk\.com/[\w\.\-]+|(?:id|club|public)\d{4,}|vkontakte\.ru/[\w\.]+', re.I)
_OK_RE  = re.compile(r'ok\.ru/[\w\.\-]+|odnoklassniki\.ru/[\w\.]+|профиль\s*ок\s*\d+', re.I)
_SOC_RE = re.compile(
    r'(?:instagram|facebook|twitter|telegram|t\.me|youtube|tiktok|linkedin)[\./][\w\.\-/]+',
    re.I,
)


def _extract_social(raw_records: list[dict]) -> tuple[str, str, str]:
    """Извлекает (vk, ok, other) из raw Sauron-записей."""
    vk_set: set[str]    = set()
    ok_set: set[str]    = set()
    other_set: set[str] = set()

    def _scan(text: str):
        for m in _VK_RE.finditer(text):
            vk_set.add(m.group(0)[:80])
        for m in _OK_RE.finditer(text):
            ok_set.add(m.group(0)[:80])
        for m in _SOC_RE.finditer(text):
            other_set.add(m.group(0)[:80])

    for rec in raw_records:
        # Поля-соцсети
        for field_name in ('ВКонтакте', 'VK', 'Vk', 'vk', 'ВК',
                           'Одноклассники', 'OK', 'Ok', 'ok',
                           'Соцсети', 'Соц.сети', 'Social'):
            v = str(rec.get(field_name) or '')
            if v.strip() and v.strip().lower() not in ('none', 'null', '-'):
                _scan(v)
        # Полный текст записи (fallback)
        for v in rec.values():
            s = str(v or '')
            if len(s) > 5:
                _scan(s)

    return (
        '; '.join(sorted(vk_set)[:5]),
        '; '.join(sorted(ok_set)[:5]),
        '; '.join(sorted(other_set)[:5]),
    )


# ═════════════════════════════════════════════════════════════════════════════
# Извлечение полей из ответа Sauron
# ═════════════════════════════════════════════════════════════════════════════

def _val(rec: dict, *keys: str) -> str:
    for k in keys:
        v = rec.get(k)
        if v and str(v).strip() and str(v).strip().lower() not in ('none', 'null', '-'):
            return str(v).strip()
    return ''


# Парсинг строки связи: «Иванов Иван Петрович (01.01.1980/муж)»
_CONN_FIO_RE = re.compile(
    r'^([А-ЯЁІЇЄ][а-яёіїє\'\-]{1,25}\s+[А-ЯЁІЇЄ][а-яёіїє\'\-]{1,20}'
    r'(?:\s+[А-ЯЁІЇЄ][а-яёіїє\'\-]{1,20})?)'
    r'(?:\s*[\(\[](.*?)[\)\]])?'
)
_DOB_RE = re.compile(r'\b(\d{1,2}[\.\/]\d{1,2}[\.\/]\d{2,4}|\d{4})\b')


def _parse_connection_str(conn: str) -> tuple[str, str, str]:
    """Из строки «Иванов ИО (01.01.1980/супруг)» → (fio, dob, relation)."""
    m = _CONN_FIO_RE.match(conn.strip())
    if not m:
        return conn.strip(), '', ''
    fio = m.group(1).strip()
    extra = (m.group(2) or '').strip()
    # DOB
    dob = ''
    dm = _DOB_RE.search(extra)
    if dm:
        dob = dm.group(1)
        extra = extra[:dm.start()] + extra[dm.end():]
    # Relation — остаток после даты
    relation = re.sub(r'[/,;]', ' ', extra).strip()
    return fio, dob, relation


def _extract_person_fields(api_result: dict) -> dict:
    """Извлекает поля из ответа Sauron. Возвращает структуру с raw_records."""
    records = api_result.get("response", [])
    if not records:
        return {
            "found_fio": "", "phones_raw": [], "emails": "",
            "addresses": "", "connections_parsed": [], "sources": "",
            "raw_count": 0, "related_fios_parsed": [], "raw_records": [],
            "vk": "", "ok": "", "other_social": "",
        }

    all_fios:     list[str] = []
    all_phones:   list[str] = []
    all_emails:   list[str] = []
    all_addrs:    list[str] = []
    all_sources:  list[str] = []
    # {conn_raw: count}
    conn_counts:  dict[str, int] = {}

    # Сбор телефонов с частотой
    phone_freq:   dict[str, int] = {}
    phone_src:    dict[str, list[str]] = {}

    for rec in records:
        fio = _val(rec, 'ФИО', 'Фамилия', 'Имя')
        if fio and fio not in all_fios:
            all_fios.append(fio)

        src = _val(rec, 'Источник', 'База', 'Источники')
        if src and src not in all_sources:
            all_sources.append(src)
        src_ctx = src  # контекст источника для этой записи

        for pf in ('Телефон', 'Телефон2', 'Телефон3', 'Phone', 'Моб', 'Mob'):
            ph = _val(rec, pf)
            if ph:
                phone_freq[ph] = phone_freq.get(ph, 0) + 1
                phone_src.setdefault(ph, []).append(src_ctx)
                if ph not in all_phones:
                    all_phones.append(ph)

        for ef in ('Email', 'E-mail', 'Эл. почта', 'Почта'):
            em = _val(rec, ef)
            if em and em not in all_emails:
                all_emails.append(em)

        addr_parts = []
        for af in ('Страна', 'Регион', 'Город', 'Населенный пункт', 'Адрес', 'Улица', 'Дом', 'Квартира'):
            v = _val(rec, af)
            if v:
                addr_parts.append(v)
        if addr_parts:
            addr = ', '.join(addr_parts)
            if addr not in all_addrs:
                all_addrs.append(addr)

        conn = _val(rec, 'Связь с лицом', 'Связанные лица', 'Связи', 'Родственники')
        if conn:
            conn_counts[conn] = conn_counts.get(conn, 0) + 1

    def join(lst: list[str], n: int = 10) -> str:
        return '; '.join(lst[:n])

    # Парсинг связанных лиц
    connections_parsed: list[dict] = []
    related_fios_parsed: list[str] = []
    for conn_raw, cnt in sorted(conn_counts.items(), key=lambda x: -x[1]):
        fio, dob, relation = _parse_connection_str(conn_raw)
        connections_parsed.append({
            'fio': fio, 'dob': dob, 'relation': relation,
            'raw': conn_raw, 'count': cnt,
        })
        if fio and len(fio.split()) >= 2 and fio not in related_fios_parsed:
            related_fios_parsed.append(fio)

    # Соцсети
    vk, ok, other_social = _extract_social(records)

    return {
        "found_fio":           join(all_fios, 3),
        "phones_raw":          all_phones,
        "phone_freq":          phone_freq,
        "phone_src":           phone_src,
        "emails":              join(all_emails, 5),
        "addresses":           join(all_addrs, 5),
        "connections_parsed":  connections_parsed,
        "sources":             join(all_sources, 10),
        "raw_count":           len(records),
        "related_fios_parsed": related_fios_parsed[:DEFAULT_MAX_RELATED],
        "raw_records":         records,
        "vk":                  vk,
        "ok":                  ok,
        "other_social":        other_social,
    }


# ═════════════════════════════════════════════════════════════════════════════
# Выбор топ-N телефонов
# ═════════════════════════════════════════════════════════════════════════════

def _pick_top_phones(
    phones_raw: list[str],
    phone_freq: dict[str, int],
    phone_src:  dict[str, list[str]],
    seen_normalized: set[str],
    top_n: int = TOP_PHONES_PER_RELATED,
    mobile_only: bool = True,
) -> tuple[list[str], int, str]:
    """
    Выбирает топ-N ликвидных телефонов по скорингу.
    Возвращает (liquid_list, discarded_count, note).
    """
    scored: list[tuple[float, str, str]] = []  # (score, norm, raw)
    seen_norm_local: set[str] = set()
    total_raw = len(phones_raw)

    for raw in phones_raw:
        norm = normalize_phone(raw)
        if norm is None:
            continue
        if norm in seen_normalized or norm in seen_norm_local:
            continue
        if mobile_only and not _is_mobile(norm):
            continue
        seen_norm_local.add(norm)
        freq = phone_freq.get(raw, 1)
        sources_ctx = ' '.join(phone_src.get(raw, []))
        score = _phone_score(raw, sources_ctx, freq)
        scored.append((score, norm, raw))

    # Сортировка по убыванию скора
    scored.sort(key=lambda x: -x[0])

    top = scored[:top_n]
    for _, norm, _ in top:
        seen_normalized.add(norm)

    liquid = [norm for _, norm, _ in top]
    discarded = total_raw - len(liquid)

    notes: list[str] = []
    if discarded > 0:
        notes.append(f"отброшено {discarded}")
    if mobile_only and any(not _is_mobile(normalize_phone(r) or '') for r in phones_raw if normalize_phone(r)):
        notes.append("городские исключены")

    return liquid, discarded, '; '.join(notes)


# ═════════════════════════════════════════════════════════════════════════════
# Пакетный поиск
# ═════════════════════════════════════════════════════════════════════════════

def batch_search(
    records: list[dict],
    chat_id: int,
    bot,
    progress_msg_id: int,
    max_records: int = DEFAULT_MAX_FIO,
    delay_sec: float = DEFAULT_DELAY_SEC,
) -> tuple[list[PersonResult], str]:
    """
    Пакетный поиск. Возвращает (results: list[PersonResult], stop_reason).
    Каждый PersonResult.related содержит список RelatedPerson (одна строка на связанного).
    """
    import sauron as _sauron

    to_search = records[:max_records]
    skipped   = len(records) - len(to_search)
    results: list[PersonResult] = []
    stop_reason = ""

    # Глобальный dedup нормализованных телефонов
    global_seen: set[str] = set()

    def _do_search(query: str) -> tuple[bool, dict, bool]:
        try:
            api_result = _sauron._api_post_search(query)
            fields = _extract_person_fields(api_result)
            stop = False
            try:
                if float(api_result.get("balance", "999")) < 1.0:
                    stop = True
            except Exception:
                pass
            return fields["raw_count"] > 0, fields, stop
        except RuntimeError as e:
            stop = any(kw in str(e) for kw in ("баланс", "Неверный", "API-ключ", "Secrets"))
            return False, {}, stop
        except Exception:
            return False, {}, False

    total = len(to_search)

    for i, record in enumerate(to_search, 1):
        fio = record['fio']

        # Прогресс
        try:
            pct = (i - 1) * 100 // total
            bar = '▓' * ((i - 1) * 10 // total) + '░' * (10 - (i - 1) * 10 // total)
            skip_note = f"\n_(пропущено по лимиту: {skipped})_" if skipped and i == 1 else ""
            bot.edit_message_text(
                f"🔍 *Поиск по файлу…*\n\n"
                f"Обработано: *{i-1}* / {total}{skip_note}\n"
                f"{bar} {pct}%\n\n"
                f"👤 Ищу: `{fio}`",
                chat_id, progress_msg_id, parse_mode="Markdown",
            )
        except Exception:
            pass

        pr = PersonResult(row_num=record['row_num'], source_fio=fio)

        # ── Основной поиск ────────────────────────────────────────────────
        found, fields, stop = _do_search(fio)

        if stop:
            pr.error = "Остановлено — недостаточно баланса или ошибка API"
            results.append(pr)
            stop_reason = pr.error
            break

        if found:
            pr.found     = True
            pr.found_fio = fields['found_fio']
            pr.emails    = fields['emails']
            pr.addresses = fields['addresses']
            pr.sources   = fields['sources']
            pr.raw_count = fields['raw_count']
            pr.vk        = fields['vk']
            pr.ok        = fields['ok']
            pr.other_social = fields['other_social']

            # Ликвидность основных телефонов
            pr.phones_raw = '; '.join(fields['phones_raw'])
            liquid, disc, note = _pick_top_phones(
                fields['phones_raw'],
                fields.get('phone_freq', {}),
                fields.get('phone_src', {}),
                global_seen,
                top_n=5,
            )
            pr.liquid_phones    = '; '.join(liquid)
            pr.discarded_phones = disc
            pr.phone_note       = note

            # Макс для основного
            pr.in_maxim, pr.maxim_source = check_maxim(
                fields['sources'], fields.get('raw_records')
            )

            # Уверенность
            if fields['found_fio']:
                q_w = set(fio.lower().split())
                f_w = set(fields['found_fio'].lower().split())
                ov  = len(q_w & f_w)
                pr.confidence = "высокая" if ov == len(q_w) else ("средняя" if ov >= 1 else "низкая")
            else:
                pr.confidence = "нет ФИО в ответе"

            # ── Связанные лица ────────────────────────────────────────────
            related_fios_parsed = fields.get('related_fios_parsed', [])
            connections_parsed  = fields.get('connections_parsed', [])

            # Строим lookup fio→metadata из connections_parsed
            conn_meta: dict[str, dict] = {}
            for cp in connections_parsed:
                conn_meta[cp['fio'].lower()] = cp

            for rel_fio in related_fios_parsed[:DEFAULT_MAX_RELATED]:
                if delay_sec > 0:
                    time.sleep(delay_sec)

                rel_found, rel_fields, rel_stop = _do_search(rel_fio)

                if rel_stop:
                    stop_reason = "Остановлено при поиске связанных — баланс или ошибка API"
                    break

                # Метаданные связи из поля connections основного человека
                meta = conn_meta.get(rel_fio.lower(), {})

                rp = RelatedPerson(
                    fio      = rel_fio,
                    dob      = meta.get('dob', ''),
                    relation = meta.get('relation', ''),
                )

                if rel_found:
                    rp.raw_count = rel_fields['raw_count']
                    rp.addresses = rel_fields['addresses']
                    rp.sources   = rel_fields['sources']
                    rp.vk        = rel_fields['vk']
                    rp.ok        = rel_fields['ok']
                    rp.other_social = rel_fields['other_social']

                    # Телефоны связанного — с приоритетом по мобильности и источникам
                    rel_liquid, rel_disc, rel_note = _pick_top_phones(
                        rel_fields['phones_raw'],
                        rel_fields.get('phone_freq', {}),
                        rel_fields.get('phone_src', {}),
                        global_seen,
                        top_n=TOP_PHONES_PER_RELATED,
                        mobile_only=True,
                    )
                    rp.liquid_phones = rel_liquid
                    rp.discarded_cnt = rel_disc
                    rp.phone_note    = rel_note

                    # Макс связанного
                    rp.in_maxim, rp.maxim_source = check_maxim(
                        rel_fields['sources'], rel_fields.get('raw_records')
                    )
                    rp.comment = f"Найдено {rp.raw_count} записей Sauron"
                else:
                    rp.comment = "Не найден в Sauron"

                pr.related.append(rp)

            if stop_reason:
                results.append(pr)
                break
        else:
            pr.confidence = "не найден"

        results.append(pr)

        if stop_reason:
            break
        if i < total:
            time.sleep(delay_sec)

    return results, stop_reason


# ═════════════════════════════════════════════════════════════════════════════
# Краткая сводка
# ═════════════════════════════════════════════════════════════════════════════

def build_short_summary(results: list[PersonResult], stop_reason: str = "") -> str:
    total     = len(results)
    found     = sum(1 for r in results if r.found)
    errors    = sum(1 for r in results if r.error and not r.found)
    not_found = total - found - errors

    rel_total  = sum(len(r.related) for r in results if r.found)
    liq_phones = sum(
        len(r.liquid_phones.split(';')) if r.liquid_phones else 0
        for r in results if r.found
    ) + sum(
        len(rp.liquid_phones)
        for r in results if r.found
        for rp in r.related
    )
    maxim_cnt = sum(
        1 for r in results
        if r.in_maxim or any(rp.in_maxim for rp in r.related)
    )
    vk_cnt = sum(1 for r in results if r.vk or any(rp.vk for rp in r.related))
    ok_cnt = sum(1 for r in results if r.ok or any(rp.ok for rp in r.related))

    lines = ["🔍 *Поиск по файлу завершён*\n"]
    lines.append(f"👤 Найдено ФИО: *{found}*")
    lines.append(f"❌ Не найдено: *{not_found}*")
    if errors:
        lines.append(f"⚠️ Ошибки: *{errors}*")
    lines.append(f"📊 Обработано: *{total}*\n")
    lines.append(f"🔗 Связанных лиц найдено: *{rel_total}*")
    lines.append(f"📱 Действующих номеров: *{liq_phones}*")
    lines.append(f"🚕 В Максе/taxsee: *{maxim_cnt}*")
    lines.append(f"🔵 ВК найдено: *{vk_cnt}*   |   🟠 ОК: *{ok_cnt}*")

    if stop_reason:
        lines.append(f"\n⛔ Остановлено: _{stop_reason}_")

    has_data = [r for r in results if r.found and (r.liquid_phones or r.related)]
    if has_data:
        lines.append("\n*Топ найденных:*")
        for r in has_data[:6]:
            ph = r.liquid_phones.split(';')[0].strip() if r.liquid_phones else ''
            mx = " 🚕" if r.in_maxim else ""
            rel_found = sum(1 for rp in r.related if rp.raw_count > 0)
            lines.append(
                f"\n👤 `{r.source_fio}`{mx}"
                + (f"\n  📱 {ph}" if ph else "")
                + (f"\n  🔗 связ.: {rel_found}" if rel_found else "")
            )
        if len(has_data) > 6:
            lines.append(f"_…ещё {len(has_data) - 6}_")

    return '\n'.join(lines)


# ═════════════════════════════════════════════════════════════════════════════
# Построение XLSX (два листа)
# ═════════════════════════════════════════════════════════════════════════════

def _make_border():
    side = Side(style='thin', color='BBBBBB')
    return Border(left=side, right=side, top=side, bottom=side)


def _write_header(ws, columns: list[tuple[str, int]],
                  header_fill, header_font, wrap):
    for col_idx, (col_name, col_width) in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font      = header_font
        cell.fill      = header_fill
        cell.alignment = wrap
        cell.border    = _make_border()
        ws.column_dimensions[cell.column_letter].width = col_width
    ws.freeze_panes = "A2"
    ws.row_dimensions[1].height = 38


def build_xlsx_report(results: list[PersonResult]) -> Optional[bytes]:
    """Два листа: «Итог по людям» и «Связанные лица»."""
    if not _HAS_OPENPYXL:
        return None

    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Итог по людям"
    ws2 = wb.create_sheet("Связанные лица")

    # ── Стили ─────────────────────────────────────────────────────────────
    H_FILL   = PatternFill("solid", fgColor="1F4E79")
    H_FILL2  = PatternFill("solid", fgColor="375623")   # зелёный заголовок для лист 2
    H_FONT   = Font(bold=True, color="FFFFFF")
    FOUND    = PatternFill("solid", fgColor="E2EFDA")
    NOTFND   = PatternFill("solid", fgColor="FCE4D6")
    MAXIM    = PatternFill("solid", fgColor="FFF2CC")
    REL_BG   = PatternFill("solid", fgColor="EEF5FB")
    REL_MAXIM= PatternFill("solid", fgColor="FFF0A0")
    WRAP     = Alignment(wrap_text=True, vertical="top")

    # ── Лист 1: «Итог по людям» ───────────────────────────────────────────
    cols1 = [
        ("Строка №",              7),
        ("Исходное ФИО",         28),
        ("Найдено",               8),
        ("Найденное ФИО",        28),
        ("Действующие тел.",     28),
        ("Email",                18),
        ("Адрес",                38),
        ("В Максе",               9),
        ("Источник Макса",       25),
        ("ВК",                   22),
        ("Одноклассники",        22),
        ("Другие соцсети",       22),
        ("Связ. лиц найдено",    10),
        ("Источники",            28),
        ("Уверенность",          12),
        ("Ошибка",               25),
    ]
    _write_header(ws1, cols1, H_FILL, H_FONT, WRAP)

    for row_idx, r in enumerate(results, 2):
        in_mx = r.in_maxim or any(rp.in_maxim for rp in r.related)
        fill  = MAXIM if in_mx and r.found else (FOUND if r.found else NOTFND)
        vals  = [
            r.row_num,
            r.source_fio,
            "Да" if r.found else "Нет",
            r.found_fio,
            r.liquid_phones,
            r.emails,
            r.addresses,
            "✅ Да" if r.in_maxim else "Нет",
            r.maxim_source,
            r.vk,
            r.ok,
            r.other_social,
            len(r.related),
            r.sources[:200] if r.sources else '',
            r.confidence,
            r.error,
        ]
        for ci, v in enumerate(vals, 1):
            cell = ws1.cell(row=row_idx, column=ci, value=v)
            cell.alignment = WRAP
            cell.border    = _make_border()
            if ci in (1, 2, 3):
                cell.fill = fill
            if ci == 8 and str(v).startswith("✅"):
                cell.fill = MAXIM
                cell.font = Font(bold=True, color="7F6000")

    # ── Лист 2: «Связанные лица» ──────────────────────────────────────────
    cols2 = [
        ("Исходное ФИО",             28),
        ("Найденное ФИО",            28),
        ("Связанное лицо ФИО",       30),
        ("Дата рождения",            14),
        ("Связь/причина",            20),
        ("Адрес связанного",         38),
        ("Действующие номера",       28),
        ("Проверка номера",          22),
        ("Есть в Максе",             10),
        ("Источник Макса",           25),
        ("ВК",                       22),
        ("Одноклассники",            22),
        ("Другие соцсети/ссылки",   22),
        ("Источники",                28),
        ("Комментарий",              25),
    ]
    _write_header(ws2, cols2, H_FILL2, H_FONT, WRAP)

    row2 = 2
    for r in results:
        if not r.found or not r.related:
            continue
        for rp in r.related:
            fill2 = REL_MAXIM if rp.in_maxim else REL_BG
            vals2 = [
                r.source_fio,
                r.found_fio,
                rp.fio,
                rp.dob,
                rp.relation,
                rp.address,
                '; '.join(rp.liquid_phones),
                rp.phone_note,
                "✅ Да" if rp.in_maxim else "Нет",
                rp.maxim_source,
                rp.vk,
                rp.ok,
                rp.other_social,
                rp.sources[:200] if rp.sources else '',
                rp.comment,
            ]
            for ci, v in enumerate(vals2, 1):
                cell = ws2.cell(row=row2, column=ci, value=v)
                cell.alignment = WRAP
                cell.border    = _make_border()
                if ci in (1, 2, 3):
                    cell.fill = fill2
                if ci == 9 and str(v).startswith("✅"):
                    cell.fill = REL_MAXIM
                    cell.font = Font(bold=True, color="7F6000")
            row2 += 1

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ═════════════════════════════════════════════════════════════════════════════
# CSV (запасной)
# ═════════════════════════════════════════════════════════════════════════════

def build_csv_report(results: list[PersonResult]) -> bytes:
    """CSV с развёрнутыми связанными лицами (одна строка = один связанный)."""
    columns = [
        "Исходное ФИО", "Строка №", "Найдено", "Найденное ФИО",
        "Действующие тел.", "Email", "Адрес", "В Максе", "Источник Макса",
        "ВК", "Одноклассники",
        "Связанное лицо ФИО", "Дата рожд.", "Связь", "Адрес связ.",
        "Номера связ.", "Проверка номера", "Связ. в Максе",
        "Ист. Макса (связ.)", "ВК связ.", "ОК связ.",
        "Источники", "Уверенность", "Ошибка",
    ]
    output = io.StringIO()
    writer = csv.writer(output, quoting=csv.QUOTE_ALL)
    writer.writerow(columns)

    for r in results:
        if r.found and r.related:
            for rp in r.related:
                writer.writerow([
                    r.source_fio, r.row_num,
                    "Да", r.found_fio,
                    r.liquid_phones, r.emails, r.addresses,
                    "Да" if r.in_maxim else "Нет", r.maxim_source,
                    r.vk, r.ok,
                    rp.fio, rp.dob, rp.relation, rp.address,
                    '; '.join(rp.liquid_phones),
                    rp.phone_note,
                    "Да" if rp.in_maxim else "Нет",
                    rp.maxim_source,
                    rp.vk, rp.ok,
                    r.sources, r.confidence, r.error,
                ])
        else:
            writer.writerow([
                r.source_fio, r.row_num,
                "Да" if r.found else "Нет", r.found_fio,
                r.liquid_phones, r.emails, r.addresses,
                "Да" if r.in_maxim else "Нет", r.maxim_source,
                r.vk, r.ok,
                '', '', '', '', '', '', '', '', '', '',
                r.sources, r.confidence, r.error,
            ])

    return output.getvalue().encode('utf-8-sig')


# ═════════════════════════════════════════════════════════════════════════════
# Preview до начала поиска
# ═════════════════════════════════════════════════════════════════════════════

def build_summary(records: list[dict], filename: str, max_records: int = DEFAULT_MAX_FIO) -> str:
    total      = len(records)
    will_check = min(total, max_records)
    skipped    = total - will_check
    lines = [f"📄 *Файл:* `{filename}`", ""]
    lines.append(f"👤 *Найдено ФИО:* {total}")
    for r in records[:8]:
        lines.append(f"  • {r['fio']}")
    if total > 8:
        lines.append(f"  _…ещё {total - 8}_")
    lines.append("")
    lines.append("━━━━━━━━━━━━━━━━━━━━")
    lines.append(f"🔍 Будет проверено: *{will_check}* ФИО")
    if skipped:
        lines.append(f"⚠️ По лимиту ({max_records}) пропущено: *{skipped}*")
    return '\n'.join(lines)


def supported_formats() -> str:
    fmts = ["txt", "csv"]
    if _HAS_OPENPYXL:
        fmts.append("xlsx")
    if _HAS_XLRD:
        fmts.append("xls")
    if _HAS_DOCX:
        fmts.append("docx")
    if _HAS_PDF:
        fmts.append("pdf")
    return ", ".join(fmts)
